"""
wordprojectconverter.py
Classe che gestisce la conversione di un xml in un documento word e il relativo salvataggio su server.
"""
import os
import re
import tempfile
import logging

from PIL import Image
from lxml import etree

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches

from .wordstyles import WordStyles
from storage.factory import StorageFactory

logger = logging.getLogger('django')


class ConverterProjectInWord:

    def __init__(self, xml_string, word_dest_path):
        self.xml_string = xml_string
        self.word_dest_path = word_dest_path

        self.xml_root = None
        self.doc = Document()
        self.styles = WordStyles(self.doc)
        self.styles.apply_styles()
        self.temp_image_files = []  # Lista per tenere traccia dei file temporanei
        self.storage = StorageFactory.get_storage_service()

    def convert_xml_string_in_node(self):
        try:
            parser = etree.XMLParser(remove_blank_text=True)
            return etree.XML(self.xml_string.encode('utf-8'), parser)
        except etree.XMLSyntaxError as e:
            logger.error(f"Errore di parsing XML: {e}")
            return None

    def download_image_for_word(self, image_path):
        """
        Scarica un'immagine da storage (locale o Spaces) e restituisce il path locale temporaneo.
        Gestisce sia lo storage locale che DigitalOcean Spaces.
        Restituisce una tupla (path_immagine, larghezza_cm, altezza_cm) o None in caso di errore.
        """
        try:
            storage_service = StorageFactory.get_storage_service()

            if hasattr(storage_service, 'download_image'):
                # Storage DigitalOcean Spaces
                try:
                    # Per DigitalOcean Spaces, il path restituito da list_files_matching_pattern
                    # è già il Key completo su S3, quindi non serve aggiungere 'media/'
                    spaces_path = image_path
                    logger.debug(f"Scaricando immagine da Spaces per Word: {spaces_path}")
                    # Scarica l'immagine come stream
                    image_stream = storage_service.download_image(spaces_path)

                    # Crea un file temporaneo
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    temp_file.write(image_stream.read())
                    temp_file.close()

                    # Aggiungi il file temporaneo alla lista per la pulizia successiva
                    self.temp_image_files.append(temp_file.name)

                    # Ottieni le dimensioni dell'immagine in cm
                    width_cm, height_cm = self.get_image_dimensions_cm(temp_file.name)

                    logger.debug(f"Immagine scaricata temporaneamente in: {temp_file.name}")
                    return temp_file.name, width_cm, height_cm

                except Exception as e:
                    return None
            else:
                # Storage locale - verifica se il file esiste
                if os.path.exists(image_path):
                    # Ottieni le dimensioni dell'immagine in cm
                    width_cm, height_cm = self.get_image_dimensions_cm(image_path)
                    return image_path, width_cm, height_cm

        except Exception as e:
            logger.error(f"Errore generale nel download dell'immagine: {e}")
            return None

    def get_image_dimensions_cm(self, image_path):
        """
        Ottiene le dimensioni dell'immagine in centimetri.
        Restituisce una tupla (larghezza_cm, altezza_cm).
        """
        try:
            with Image.open(image_path) as img:
                # Ottieni le dimensioni in pixel
                width_px, height_px = img.size
                
                # Converti in cm (assumendo 96 DPI come standard)
                # 1 pollice = 2.54 cm, 96 DPI = 96 pixel per pollice
                dpi = 96
                width_cm = (width_px / dpi) * 2.54
                height_cm = (height_px / dpi) * 2.54
                
                return width_cm, height_cm
        except Exception as e:
            logger.error(f"Errore nel calcolo delle dimensioni dell'immagine {image_path}: {e}")
            return None, None

    def cleanup_temp_files(self):
        """
        Pulisce tutti i file temporanei creati per le immagini.
        """
        for temp_file in self.temp_image_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
                    logger.debug(f"File temporaneo eliminato: {temp_file}")
            except Exception as e:
                logger.error(f"Errore nell'eliminazione del file temporaneo {temp_file}: {e}")
        self.temp_image_files = []

    def add_metadata(self):
        """Aggiungi metadati come titolo, autore, data, ecc."""
        metadata = self.xml_root.find("metadata")

        if metadata is not None:
            title = metadata.find("title").text if metadata.find("title") is not None else ""
            author = metadata.find("author").text if metadata.find("author") is not None else ""
            version = metadata.find("version").text if metadata.find("version") is not None else ""
            publisher = metadata.find("publisher").text if metadata.find("publisher") is not None else ""
            comments = metadata.find("comments").text if metadata.find("comments") is not None else ""

            self.doc.core_properties.title = title
            self.doc.core_properties.author = author
            self.doc.core_properties.version = version
            self.doc.core_properties.publisher = publisher
            self.doc.core_properties.comments = comments
            self.doc.core_properties.language = "it-IT"

    def add_header(self):
        """Aggiungi l'intestazione alla sezione corrente (inizia dalla pagina 4)."""
        section = self.doc.sections[-1]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Contenuto dell'intestazione"

    def add_footer(self, section):
        """Aggiunge un piè di pagina centrato con il numero di pagina."""
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centra il numero di pagina

        # Aggiunge il campo per il numero di pagina nel piè di pagina
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE')  # Campo per il numero di pagina

        run = paragraph.add_run()
        run._r.append(fld_simple)  # Aggiunge il campo come run

    '''
    def add_header(self, start_page=4):
        """Aggiungi l'intestazione dal XML."""
        header = self.xml_root.find("header")
        if header is not None:
            header_text = header.find("text").text if header.find(
                "text") is not None else ""
            section = self.doc.sections[0]
            header_paragraph = section.header.paragraphs[0]
            header_paragraph.text = header_text

    def add_footer(self, start_page=4):
        """Aggiungi il piè di pagina dal XML."""
        footer = self.xml_root.find("footer")
        if footer is not None:
            footer_text = footer.find("text").text if footer.find(
                "text") is not None else ""
            section = self.doc.sections[0]
            footer_paragraph = section.footer.paragraphs[0]
            footer_paragraph.text = footer_text
    '''

    def add_summary(self):
        """Aggiungi il sommario."""
        # Aggiungi un segnaposto per il sommario
        self.doc.add_paragraph('Indice', style='Heading 1')

        summary = self.xml_root.find('summary')
        count = 1
        for elem in summary:

            if elem.tag == "subtitle":
                if elem.text:
                    cleaned_text = re.sub(r"^\d+\.\s*", "", elem.text.strip())
                    self.doc.add_paragraph("")
                    self.doc.add_heading(cleaned_text, level=3)
            elif elem.tag == "text":
                text_parts = []
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                for bold in elem.findall("bold"):
                    if bold.text:
                        text_parts.append(bold.text.strip())
                    if bold.tail and bold.tail.strip():
                        text_parts.append(bold.tail.strip())
                full_text = " ".join(text_parts)
                self.doc.add_paragraph(full_text, style='Normal')
                count += 1
            elif elem.tag == "empty-line":
                self.doc.add_paragraph("")
            elif elem.tag == "add-page-break":
                self.doc.add_page_break()

        self.doc.add_page_break()

    def add_chapters(self):
        """Aggiungi capitoli e paragrafi dal XML."""

        chapter_counter = 0

        for chapter in self.xml_root.find("chapters").findall("chapter"):
            chapter_counter += 1
            paragraph_counter = 1
            figure_counter = 1
            table_counter = 1

            # Aggiungi un'interruzione di pagina all'inizio di ogni capitolo
            if self.doc.paragraphs and chapter_counter > 1:
                self.doc.add_page_break()

            for elem in chapter:
                if elem.tag == "h1":
                    chapter_title = chapter.find("h1").text if chapter.find("h1") is not None else ""
                    self.doc.add_heading(f"{chapter_counter}. {chapter_title}", level=1)
                elif elem.tag == "title":
                    if chapter_counter != 4 and (paragraph_counter != 1 or paragraph_counter != 2):
                        self.doc.add_paragraph("")
                    self.doc.add_heading(f"{chapter_counter}.{paragraph_counter} {elem.text.strip()}", level=2)
                    paragraph_counter += 1  # Incrementa il contatore dei paragrafi
                elif elem.tag == "subtitle":
                    if elem.text:
                        cleaned_text = re.sub(r"^\d+\.\s*", "", elem.text.strip())
                        self.doc.add_paragraph("")
                        self.doc.add_heading(cleaned_text, level=3)
                elif elem.tag == "text":
                    text_parts = []
                    if elem.text and elem.text.strip():
                        text_parts.append(elem.text.strip())
                    for bold in elem.findall("bold"):
                        if bold.text:
                            text_parts.append(bold.text.strip())
                        if bold.tail and bold.tail.strip():
                            text_parts.append(bold.tail.strip())
                    full_text = " ".join(text_parts)
                    self.doc.add_paragraph(full_text, style='Normal')
                elif elem.tag == "empty-line":
                    self.doc.add_paragraph("")
                elif elem.tag == "add-page-break":
                    self.doc.add_page_break()
                elif elem.tag == "italic":
                    run = self.doc.add_paragraph().add_run(elem.text)
                    run.italic = True
                elif elem.tag == "ul" or elem.tag == "ol":
                    list_style = 'List Bullet' if elem.tag == "ul" else 'List Bullet'  # Number'
                    for li in elem.findall("li"):
                        full_text = []
                        # Testo scritto direttamente dentro li senza altri tag
                        if li.text and li.text.strip():
                            full_text.append(li.text.strip())
                        # Itera su tutti gli elementi dentro <li>
                        for sub_elem in li:
                            if sub_elem.tag in ["text", "bold", "italic"] and sub_elem.text:
                                full_text.append(sub_elem.text)
                            elif sub_elem.text:
                                full_text.append(sub_elem.text)
                            # Aggiungi anche il testo dopo ogni tag (tail)
                            if sub_elem.tail:
                                full_text.append(sub_elem.tail)
                        final_text = " ".join(full_text).strip()
                        if final_text:
                            self.doc.add_paragraph(
                                final_text, style=list_style)
                elif elem.tag == "table":
                    word_table = None
                    caption_element = elem.find("caption")
                    if caption_element is not None:
                        self.add_empty_paragraphs(1)
                        self.doc.add_paragraph(f"{chapter_counter}.{table_counter} - {caption_element.text}",
                                               style='Caption')
                        table_counter += 1
                    rows = elem.findall("row")
                    if rows:
                        num_cols = len(rows[0])
                        word_table = self.doc.add_table(
                            rows=len(rows), cols=num_cols)
                        for i, row in enumerate(rows):
                            cells = row.findall("cell")
                            for j, cell in enumerate(cells):
                                text_parts = []
                                # Se <cell> ha del testo diretto, lo usa
                                if cell.text and cell.text.strip():
                                    text_parts.append(cell.text.strip())
                                # Se c'è un tag <bold>
                                bold_element = cell.find("bold")
                                if bold_element is not None:
                                    text_parts.append(
                                        bold_element.text.strip())
                                # Se c'è un tag <text> dentro <cell>, lo usa ignorando <bold>
                                text_element = cell.find("text")
                                if text_element is not None:
                                    text_content = []
                                    # Se <text> ha del testo diretto, lo prende
                                    if text_element.text and text_element.text.strip():
                                        text_content.append(
                                            text_element.text.strip())
                                    # Se <text> ha <bold>, prende solo il contenuto ignorando il tag
                                    for bold in text_element.findall("bold"):
                                        if bold.text and bold.text.strip():
                                            text_content.append(
                                                bold.text.strip())
                                        if bold.tail and bold.tail.strip():
                                            text_content.append(
                                                bold.tail.strip())
                                    # Unisce il testo dentro <text> e lo aggiunge ai text_parts
                                    text_parts.append(" ".join(text_content))
                                # Combina il testo e lo inserisce nella cella della tabella Word
                                final_text = " ".join(text_parts).strip()
                                word_table.cell(i, j).text = final_text
                    self.styles.apply_table_style(
                        word_table, elem.get('style'))
                elif elem.tag == "image":
                    image_path = elem.get('src')
                    image_style = elem.get('style')
                    image_width = elem.get('width')  # L'attributo width in cm

                    if image_path:
                        # Scarica l'immagine temporaneamente se necessario
                        result = self.download_image_for_word(image_path)
                        if result:
                            local_image_path, actual_width_cm, actual_height_cm = result
                            
                            image_width_in_inches = None
                            if image_width:
                                image_width_in_inches = float(image_width) / 2.54

                            image_paragraph = self.doc.add_paragraph()
                            image_run = image_paragraph.add_run()

                            image_run.add_picture(local_image_path,
                                                  width=Inches(image_width_in_inches) if image_width_in_inches else None)

                            if image_style == "center":
                                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            caption_element = elem.find("caption")
                            if caption_element is not None:
                                self.doc.add_paragraph(f"{chapter_counter}.{figure_counter} - {caption_element.text}",
                                                       style='Caption')
                                figure_counter += 1

    def add_empty_paragraphs(self, n_empty_paragraphs):
        """Aggiunge un numero definito di paragrafi vuoti al documento"""

        for _ in range(n_empty_paragraphs):
            self.doc.add_paragraph("", style='Normal')

    def add_cover_page(self):
        """Aggiungi la copertina di progetto."""

        cover_page = self.xml_root.find("cover-page")

        if cover_page is not None:
            # Loghi
            table_logos = cover_page.find("table[@position='logos']")
            if table_logos is not None:
                self.add_styled_table(table_logos)
                self.add_empty_paragraphs(1)

            # Titolo
            title = cover_page.find("title").text if cover_page.find("title") is not None else ""
            self.add_styled_paragraph(title, 'CoverH1', False, 1)
            self.add_empty_paragraphs(1)

            # Eventuale immagine dopo il titolo
            image_after_title = cover_page.find("image[@position='after_title']")
            if image_after_title is not None:
                image_path = image_after_title.get("src")
                if image_path:
                    image_data = self.storage.download_image(image_path)
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.add_run().add_picture(image_data)
            self.add_empty_paragraphs(1)

            # Sottotitolo
            subtitle = cover_page.find("subtitle")
            if subtitle is not None:
                for text_element in subtitle.findall("text"):
                    self.add_styled_paragraph(text_element.text, style='CoverH2', upper=False, alignment=1)
            self.add_empty_paragraphs(1)

            # Tipo di documento
            doc_types = cover_page.findall("doc-type")
            for doc_type in doc_types:
                self.add_styled_paragraph(doc_type.text, 'Normal', upper=True, alignment=1)
            self.add_empty_paragraphs(1)

            # Verifica e aggiunge tabella del team di progetto
            table_project_team = cover_page.find("table[@position='project_team']")
            self.add_styled_table(table_project_team)

            # Verifica e aggiunge tabella della revisione
            table_revision = cover_page.find("table[@position='revision']")
            self.add_styled_table(table_revision)
            # self.add_empty_paragraphs(1)

            table_revision = cover_page.find("table[@position='file-code']")
            self.add_styled_table(table_revision)

    def add_footer_cover_page(self, section):
        footer_cover_page = self.xml_root.find("footer-cover-page")
        if footer_cover_page is not None:
            # Aggiungi il piè di pagina specifico per la prima pagina
            first_page_footer = section.footer
            # paragraph = first_page_footer.paragraphs[0]
            # paragraph.text = "Piè di pagina solo per la prima pagina"

            # Verifica e aggiunge tabella dopo i tipi di documento, se presente
            table_revision = footer_cover_page.find("table[@position='revision']")
            if table_revision is not None:
                word_table = None
                rows = table_revision.findall("row")
                if rows:
                    num_cols = len(rows[0])
                    word_table = first_page_footer.add_table(
                        rows=len(rows), cols=num_cols, width=Inches(6))
                    for i, row in enumerate(rows):
                        cells = row.findall("cell")
                        for j, cell in enumerate(cells):
                            # Usa uno spazio vuoto come valore predefinito
                            cell_text = cell.text if cell.text is not None else " "
                            word_table.cell(i, j).text = cell_text

                self.styles.apply_table_style(
                    word_table,  table_revision.get('style'))

    def add_styled_paragraph(self, text, style, upper=False, alignment=2):
        """Aggiungi un paragrafo secondo lo stile indicato."""

        if text:
            if upper:
                text = text.upper()
            para = self.doc.add_paragraph(text, style=style)
            para.alignment = alignment

    def add_styled_table(self, table_node):
        if table_node is not None:
            word_table = None
            rows = table_node.findall("row")
            if rows:
                num_cols = len(rows[0])
                word_table = self.doc.add_table(rows=len(rows), cols=num_cols)
                for i, row in enumerate(rows):
                    cells = row.findall("cell")
                    for j, cell in enumerate(cells):
                        # Testo
                        # Usa uno spazio vuoto come valore predefinito
                        cell_text = cell.text if cell.text is not None else " "
                        word_table.cell(i, j).text = cell_text
                        # Immagine
                        image_node = cell.find("image")
                        if image_node is not None:
                            image_path = image_node.get("src")
                            if image_path:
                                image_data = self.storage.download_image(image_path)
                                paragraph = word_table.cell(i, j).paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(image_data)
                                paragraph.alignment = 1

            self.styles.apply_table_style(word_table, table_node.get('style'))

    def add_back_cover(self):
        """Aggiungi la pagina di retro copertina."""

        back_cover = self.xml_root.find("back-cover")

        if back_cover is not None:
            # Titolo
            title = back_cover.find("title").text if back_cover.find("title") is not None else ""
            self.add_styled_paragraph(title, 'CoverH2', upper=False, alignment=1)
            self.add_empty_paragraphs(2)

            # Sottotitolo
            subtitle = back_cover.find("subtitle")
            if subtitle is not None:
                for text_element in subtitle.findall("text"):
                    self.add_styled_paragraph(text_element.text, 'CoverH2', upper=False, alignment=1)
            self.add_empty_paragraphs(1)

            # Tipo di documento
            doc_types = back_cover.findall("doc-type")
            for doc_type in doc_types:
                self.add_styled_paragraph(doc_type.text, 'Normal', upper=True, alignment=1)
            self.add_empty_paragraphs(1)

            # Tabella committente e proponente
            table = back_cover.find("table[@position='client']")
            self.add_styled_table(table)

            # Tabella progettazione tecnica
            table = back_cover.find("table[@position='designer']")
            self.add_styled_table(table)
            self.add_empty_paragraphs(1)

    def add_final_cover(self):
        """Aggiungi la pagina di retro copertina."""

        back_cover = self.xml_root.find("final-cover")
        if back_cover is not None:
            # Aggiungi un'interruzione di pagina dopo il back cover
            self.doc.add_page_break()

            legal_info = back_cover.find("legal-information").text if back_cover.find("legal-information") is not None else ""
            self.doc.add_paragraph(legal_info, style='Normal')

    def create_document(self):
        """Crea il documento Word a partire dall'XML."""

        try:
            self.add_metadata()
            self.add_cover_page()

            # Aggiungi il footer specifico solo per la copertina (contiene due tabelle revisione)
            first_section = self.doc.sections[0]
            # Abilita un piè di pagina specifico per la prima pagina
            first_section.footer.is_linked_to_previous = False
            first_section.different_first_page_header_footer = True
            self.add_footer_cover_page(first_section)
            self.doc.add_page_break()

            # Aggiungi una interruzione di sezione
            # section = self.doc.add_section(WD_SECTION.NEW_PAGE)
            # section.header.is_linked_to_previous = False
            # section.footer.is_linked_to_previous = False

            self.add_back_cover()

            # Sezione principale
            main_section = self.doc.add_section(WD_SECTION.NEW_PAGE)

            self.add_summary()

            # Inizia la numerazione della pagina in questa sezione a partire da 1
            main_section.start_page_number = 1  # QUESTO NON FUNZIONA

            # Disabilita il collegamento per le intestazioni e i piè di pagina
            main_section.header.is_linked_to_previous = False
            main_section.footer.is_linked_to_previous = False

            # Aggiungi header e footer alla quarta pagina
            # self.add_header()
            self.add_footer(main_section)

            # Aggiungi i capitoli
            self.add_chapters()

            # Aggiungi la final cover
            self.add_final_cover()

        finally:
            # Pulisci sempre i file temporanei alla fine
            self.cleanup_temp_files()

    def save_document(self):
        """Salva il documento Word in un file temporaneo e lo carica nello storage."""
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = temp_file.name
                self.doc.save(temp_path)

            destination_path = self.storage.upload_file(temp_path, self.word_dest_path)
            logger.info(f"Documento salvato con successo in {destination_path}")
            return destination_path
        except Exception as e:
            logger.error(f"Errore nel salvataggio del documento nello storage: {e}", exc_info=True)
            return False
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    def create_and_save_word(self):
        """Gestisce l'intero processo: parsing XML, generazione Word, salvataggio."""

        # Converti stringa xml in root xml
        self.xml_root = self.convert_xml_string_in_node()
        if self.xml_root is None:
            logger.warning("Interrotto: XML non valido, impossibile generare il file Word.")
            return False

        # Crea il word
        self.create_document()

        # Salva il documento word
        return self.save_document()
