"""
wordstyles.py
Classe che include e gestisce tutti gli stili applicati al documenti word.
"""
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class WordStyles:
    def __init__(self, doc):
        self.doc = doc

    def set_document_language(self, language_code='it-IT'):
        """Imposta la lingua del documento modificando il nodo <w:lang> in /word/styles.xml."""

        styles = self.doc.styles
        doc_defaults = styles.element.find(qn('w:docDefaults'))

        if doc_defaults is None:
            # Se non esiste <w:docDefaults>, creiamo tutta la struttura necessaria
            doc_defaults = OxmlElement('w:docDefaults')
            self.doc.styles.element.append(doc_defaults)

        rpr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rpr_default is None:
            rpr_default = OxmlElement('w:rPrDefault')
            doc_defaults.append(rpr_default)

        rpr = rpr_default.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            rpr_default.append(rpr)

        lang_element = rpr.find(qn('w:lang'))
        if lang_element is None:
            lang_element = OxmlElement('w:lang')
            rpr.append(lang_element)

        lang_element.set(qn('w:val'), language_code)

    def apply_styles(self):
        """Applica gli stili e le personalizzazioni al documento Word."""

        # Imposta la lingua del documento
        self.set_document_language('it-IT')

        # Margini del documento in centimetri
        section = self.doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Personalizza lo stile 'Normal'
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        normal_font.color.rgb = RGBColor(0, 0, 0)

        paragraph_format = normal_style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1  # Imposta l'interlinea a singola

        # Personalizza lo stile 'Heading 1'
        heading1_style = self.doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Calibri'
        heading1_font.size = Pt(18)
        heading1_font.bold = False
        heading1_font.color.rgb = RGBColor(0, 0, 0)

        # Personalizza lo stile 'Heading 2'
        heading2_style = self.doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Calibri'
        heading2_font.size = Pt(14)
        heading2_font.bold = False
        heading2_font.color.rgb = RGBColor(0, 0, 0)

        # Aggiungi spazio prima del titolo Heading 2
        heading2_paragraph_format = heading2_style.paragraph_format
        heading2_paragraph_format.space_before = Pt(18)
        heading2_paragraph_format.line_spacing = 1  # Imposta l'interlinea a singola

        # Personalizza lo stile 'Heading 3'
        heading3_style = self.doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = 'Calibri'
        heading3_font.size = Pt(11)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(0, 0, 0)
        heading3_paragraph_format = heading3_style.paragraph_format
        heading3_paragraph_format.space_after = Pt(6)
        heading3_paragraph_format.line_spacing = 1  # Imposta l'interlinea a singola

        # Personalizza lo stile "Caption"
        caption_style = self.doc.styles['Caption']
        caption_font = caption_style.font
        caption_font.name = 'Calibri'
        caption_font.size = Pt(10)
        caption_font.italic = False
        caption_font.bold = False
        caption_font.color.rgb = RGBColor(0, 0, 0)
        caption_format = caption_style.paragraph_format
        caption_format.space_before = Pt(6)
        caption_format.space_after = Pt(6)

        # Personalizza lo stile "List Bullet"
        list_style = self.doc.styles['List Bullet']
        list_font = list_style.font
        list_font.name = 'Calibri'
        list_font.size = Pt(11)
        list_font.color.rgb = RGBColor(0, 0, 0)  # Colore del testo nella lista: Nero (RGB)

        # Personalizza lo stile "List Number"
        list_style = self.doc.styles['List Number']
        list_font = list_style.font
        list_font.name = 'Calibri'
        list_font.size = Pt(11)
        list_font.color.rgb = RGBColor(0, 0, 0)  # Colore del testo nella lista: Nero (RGB)

        # Stile personalizzato per il sommario
        if 'TOC Heading' not in self.doc.styles:
            toc_style = self.doc.styles.add_style('TOC Heading', 1)  # Stile paragrafo
            toc_font = toc_style.font
            toc_font.name = 'Calibri'
            toc_font.size = Pt(12)
            toc_font.color.rgb = RGBColor(0, 0, 0)

        # Personalizza lo stile in copertina H0
        if 'CoverH0' not in self.doc.styles:
            cover_h0_style = self.doc.styles.add_style('CoverH0', 1)  # 1 indica che è uno stile di paragrafo
        else:
            cover_h0_style = self.doc.styles['CoverH0']
        cover_h0_font = cover_h0_style.font
        cover_h0_font.name = 'Calibri'
        cover_h0_font.size = Pt(48)
        cover_h0_font.bold = False
        cover_h0_font.color.rgb = RGBColor(0, 0, 0)
        cover_h0_paragraph_format = cover_h0_style.paragraph_format
        cover_h0_paragraph_format.space_before = Pt(6)
        cover_h0_paragraph_format.space_after = Pt(6)
        cover_h0_paragraph_format.line_spacing = 1

        # Personalizza lo stile in copertina H1
        if 'CoverH1' not in self.doc.styles:
            cover_h1_style = self.doc.styles.add_style('CoverH1', 1)  # 1 indica che è uno stile di paragrafo
        else:
            cover_h1_style = self.doc.styles['CoverH1']
        cover_h1_font = cover_h1_style.font
        cover_h1_font.name = 'Calibri'
        cover_h1_font.size = Pt(18)
        cover_h1_font.bold = False
        cover_h1_font.color.rgb = RGBColor(0, 0, 0)
        cover_h1_paragraph_format = cover_h1_style.paragraph_format
        cover_h1_paragraph_format.space_before = Pt(6)
        cover_h1_paragraph_format.space_after = Pt(6)
        cover_h1_paragraph_format.line_spacing = 1

        # Personalizza lo stile in copertina H2
        if 'CoverH2' not in self.doc.styles:
            cover_h2_style = self.doc.styles.add_style('CoverH2', 1)  # 1 indica che è uno stile di paragrafo
        else:
            cover_h2_style = self.doc.styles['CoverH2']
        cover_h2_font = cover_h2_style.font
        cover_h2_font.name = 'Calibri'
        cover_h2_font.size = Pt(14)
        cover_h2_font.bold = False
        cover_h2_font.color.rgb = RGBColor(0, 0, 0)
        cover_h2_paragraph_format = cover_h2_style.paragraph_format
        cover_h2_paragraph_format.space_before = Pt(0)
        cover_h2_paragraph_format.space_after = Pt(0)
        cover_h2_paragraph_format.line_spacing = 1

    def apply_table_style(self, word_table, style):
        """
        Applica uno stile a una tabella (bordi, allineamenti, dimensione e stile del font, ...).
        borders:TRBL;
        width:10cm;alignment:center;font-size:12; padding=3
        column-alignments:L-C-C-R; width-cols:3cm-2cm-2cm-3cm;font-style:B-B-I-N;
        """

        if word_table is None:
            return

        style_dict = self._parse_style(style)

        # Larghezza totale della tabella in cm
        if "width" in style_dict:
            total_width = float(style_dict["width"])
        else:
            total_width = 17

        # Larghezza delle singole colonne
        if "width-cols" in style_dict:
            fixed_widths = self._parse_column_widths(style_dict.get("width-cols", ""))
            num_columns = len(word_table.columns)

            if len(fixed_widths) == num_columns:
                for idx, column in enumerate(word_table.columns):
                    column_width = fixed_widths[idx]
                    for cell in column.cells:
                        cell.width = Cm(column_width)
            else:
                num_fixed_columns = len(fixed_widths)
                fixed_width = sum(fixed_widths)
                remaining_width = total_width - fixed_width
                dynamic_width = round(remaining_width / (num_columns - num_fixed_columns), 1)

                for idx, column in enumerate(word_table.columns):
                    for cell in column.cells:
                        if idx < num_fixed_columns:
                            cell.width = Cm(fixed_widths[idx])
                        else:
                            cell.width = Cm(dynamic_width)

        # Bordi personalizzati
        if "borders" in style_dict:
            self._apply_borders(word_table, style_dict["borders"])

        # Allineamento della tabella
        if "alignment" in style_dict:
            alignment = style_dict["alignment"]
            if alignment == "center":
                word_table.alignment = 1
            elif alignment == "left":
                word_table.alignment = 0
            elif alignment == "right":
                word_table.alignment = 2

        # Stile del testo (font_size, font_style)
        font_size = int(style_dict.get("font-size", 11))
        font_styles = style_dict.get("font-style", "").split('-')

        # Padding
        padding = int(style_dict.get("padding", 3))

        # Allineamento dentro le colonne
        if "column-alignments" in style_dict:
            column_alignments = style_dict["column-alignments"].split("-")
        else:
            column_alignments = []

        if "space-before" in style_dict:
            self.doc.add_paragraph()
            para = self.doc.paragraphs[-1]
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)

        for row in word_table.rows:
            for idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:

                    col_alignment = column_alignments[idx] if idx < len(column_alignments) else ''
                    if col_alignment == "L":
                        cell.paragraphs[0].alignment = 0
                    elif col_alignment == "C":
                        cell.paragraphs[0].alignment = 1
                    elif col_alignment == "R":
                        cell.paragraphs[0].alignment = 2

                    paragraph.paragraph_format.space_before = Pt(padding)
                    paragraph.paragraph_format.space_after = Pt(padding)

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Calibri'

                        # Stile del font in base alla posizione della colonna
                        if idx < len(font_styles):
                            style = font_styles[idx]

                            if "B" in style:
                                run.font.bold = True
                            if "I" in style:
                                run.font.italic = True
                            if "BI" in style:
                                run.font.bold = True
                                run.font.italic = True
                            if "N" in style:
                                run.font.bold = False
                                run.font.italic = False
        if "space-after" in style_dict:
            self.doc.add_paragraph()

    @staticmethod
    def _parse_column_widths(width_cols):
        """Converte la stringa delle larghezze delle colonne (es: "3-2-2-3") in una lista di numeri."""

        widths = []
        for width in width_cols.split('-'):
            widths.append(float(width))
        return widths

    @staticmethod
    def _parse_style(style):
        """Converte la stringa di stile in un dizionario."""

        style_dict = {}
        if style:
            for item in style.split(";"):
                if ":" in item:
                    key, value = item.split(":")
                    style_dict[key.strip()] = value.strip()
        return style_dict

    @staticmethod
    def _apply_borders(word_table, borders):
        """Applica bordi personalizzati alle celle della tabella."""

        if word_table is not None:
            border_map = {"T": "top", "B": "bottom", "L": "left", "R": "right"}
            for row in word_table.rows:
                for cell in row.cells:
                    cell_borders = cell._element.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border in borders:
                        if border in border_map:
                            side = border_map[border]
                            border = OxmlElement(f'w:{side}')
                            border.set(qn('w:val'), 'single')           # Tipo di bordo
                            border.set(qn('w:sz'), '5')                # Dimensione del bordo (in ottavi di punto)
                            border.set(qn('w:space'), '0')              # Spaziatura opzionale
                            border.set(qn('w:color'), '000000')         # Colore del bordo
                            tcBorders.append(border)

                            if cell_borders.find(qn('w:tcBorders')) is not None:
                                cell_borders.remove(cell_borders.find(qn('w:tcBorders')))
                            cell_borders.append(tcBorders)
