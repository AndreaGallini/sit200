"""
worddesignconverter.py
Classe che gestisce la conversione di un xml in un documento word e il relativo salvataggio su server.
"""
import os
import re
import tempfile
import logging

from PIL import Image
from lxml import etree

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches

from .wordstyles import WordStyles
from storage.factory import StorageFactory

logger = logging.getLogger('django')


class ConverterDesignInWord:

    def __init__(self, xml_string, word_dest_path):
        self.xml_string = xml_string
        self.word_dest_path = word_dest_path

        self.xml_root = None
        self.doc = Document()
        self.styles = WordStyles(self.doc)
        self.styles.apply_styles()
        self.storage = StorageFactory.get_storage_service()

    def convert_xml_string_in_node(self):
        try:
            parser = etree.XMLParser(remove_blank_text=True)
            return etree.XML(self.xml_string.encode('utf-8'), parser)
        except etree.XMLSyntaxError as e:
            logger.error(f"Errore di parsing XML: {e}")
            return None

    def add_metadata(self):
        """Aggiungi metadati come titolo, autore, data, ecc."""
        metadata = self.xml_root.find("metadata")

        if metadata is not None:
            title = metadata.find("title").text if metadata.find("title") is not None else ""
            author = metadata.find("author").text if metadata.find("author") is not None else ""
            version = metadata.find("version").text if metadata.find("version") is not None else ""
            publisher = metadata.find("publisher").text if metadata.find("publisher") is not None else ""
            comments = metadata.find("comments").text if metadata.find("comments") is not None else ""

            self.doc.core_properties.title = title
            self.doc.core_properties.author = author
            self.doc.core_properties.version = version
            self.doc.core_properties.publisher = publisher
            self.doc.core_properties.comments = comments
            self.doc.core_properties.language = "it-IT"

    def add_header(self):
        """Aggiungi l'intestazione alla sezione corrente (inizia dalla pagina 4)."""
        section = self.doc.sections[-1]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Contenuto dell'intestazione"

    def add_footer(self, section):
        """Aggiunge un piè di pagina centrato con il numero di pagina."""
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centra il numero di pagina

        # Aggiunge il campo per il numero di pagina nel piè di pagina
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE')  # Campo per il numero di pagina

        run = paragraph.add_run()
        run._r.append(fld_simple)  # Aggiunge il campo come run

    def add_chapters(self):
        """Aggiungi capitoli e paragrafi dal XML."""

        chapter_counter = 0

        for chapter in self.xml_root.find("chapters").findall("chapter"):
            chapter_counter += 1
            paragraph_counter = 1
            figure_counter = 1
            table_counter = 1

            # Aggiungi un'interruzione di pagina all'inizio di ogni capitolo
            if self.doc.paragraphs and chapter_counter > 1:
                self.doc.add_page_break()

            for elem in chapter:
                if elem.tag == "h1":
                    chapter_title = chapter.find(
                        "h1").text if chapter.find("h1") is not None else ""
                    self.doc.add_heading(
                        f"{chapter_counter}. {chapter_title}", level=1)
                elif elem.tag == "title":
                    if chapter_counter != 4 and (paragraph_counter != 1 or paragraph_counter != 2):
                        self.doc.add_paragraph("")
                    self.doc.add_heading(
                        f"{chapter_counter}.{paragraph_counter} {elem.text.strip()}", level=2)
                    paragraph_counter += 1  # Incrementa il contatore dei paragrafi
                elif elem.tag == "subtitle":
                    if elem.text:
                        cleaned_text = re.sub(
                            r"^\d+\.\s*", "", elem.text.strip())
                        self.doc.add_paragraph("")
                        self.doc.add_heading(cleaned_text, level=3)
                elif elem.tag == "text":
                    text_parts = []
                    if elem.text and elem.text.strip():
                        text_parts.append(elem.text.strip())
                    for bold in elem.findall("bold"):
                        if bold.text:
                            text_parts.append(bold.text.strip())
                        if bold.tail and bold.tail.strip():
                            text_parts.append(bold.tail.strip())
                    full_text = " ".join(text_parts)
                    self.doc.add_paragraph(full_text, style='Normal')
                elif elem.tag == "empty-line":
                    self.doc.add_paragraph("")
                elif elem.tag == "add-page-break":
                    self.doc.add_page_break()
                elif elem.tag == "italic":
                    run = self.doc.add_paragraph().add_run(elem.text)
                    run.italic = True
                elif elem.tag == "ul" or elem.tag == "ol":
                    list_style = 'List Bullet' if elem.tag == "ul" else 'List Bullet'  # Number'
                    for li in elem.findall("li"):
                        full_text = []
                        # Testo scritto direttamente dentro li senza altri tag
                        if li.text and li.text.strip():
                            full_text.append(li.text.strip())
                        # Itera su tutti gli elementi dentro <li>
                        for sub_elem in li:
                            if sub_elem.tag in ["text", "bold", "italic"] and sub_elem.text:
                                full_text.append(sub_elem.text)
                            elif sub_elem.text:
                                full_text.append(sub_elem.text)
                            # Aggiungi anche il testo dopo ogni tag (tail)
                            if sub_elem.tail:
                                full_text.append(sub_elem.tail)
                        final_text = " ".join(full_text).strip()
                        if final_text:
                            self.doc.add_paragraph(
                                final_text, style=list_style)
                elif elem.tag == "table":
                    word_table = None
                    caption_element = elem.find("caption")
                    if caption_element is not None:
                        self.add_empty_paragraphs(1)
                        self.doc.add_paragraph(f"{chapter_counter}.{table_counter} - {caption_element.text}",
                                               style='Caption')
                        table_counter += 1
                    rows = elem.findall("row")
                    if rows:
                        num_cols = len(rows[0])
                        word_table = self.doc.add_table(
                            rows=len(rows), cols=num_cols)
                        for i, row in enumerate(rows):
                            cells = row.findall("cell")
                            for j, cell in enumerate(cells):
                                text_parts = []
                                # Se <cell> ha del testo diretto, lo usa
                                if cell.text and cell.text.strip():
                                    text_parts.append(cell.text.strip())
                                # Se c'è un tag <bold>
                                bold_element = cell.find("bold")
                                if bold_element is not None:
                                    text_parts.append(
                                        bold_element.text.strip())
                                # Se c'è un tag <text> dentro <cell>, lo usa ignorando <bold>
                                text_element = cell.find("text")
                                if text_element is not None:
                                    text_content = []
                                    # Se <text> ha del testo diretto, lo prende
                                    if text_element.text and text_element.text.strip():
                                        text_content.append(
                                            text_element.text.strip())
                                    # Se <text> ha <bold>, prende solo il contenuto ignorando il tag
                                    for bold in text_element.findall("bold"):
                                        if bold.text and bold.text.strip():
                                            text_content.append(
                                                bold.text.strip())
                                        if bold.tail and bold.tail.strip():
                                            text_content.append(
                                                bold.tail.strip())
                                    # Unisce il testo dentro <text> e lo aggiunge ai text_parts
                                    text_parts.append(" ".join(text_content))
                                # Combina il testo e lo inserisce nella cella della tabella Word
                                final_text = " ".join(text_parts).strip()
                                word_table.cell(i, j).text = final_text
                    self.styles.apply_table_style(
                        word_table, elem.get('style'))

    def add_empty_paragraphs(self, n_empty_paragraphs):
        """Aggiunge un numero definito di paragrafi vuoti al documento"""

        for _ in range(n_empty_paragraphs):
            self.doc.add_paragraph("", style='Normal')

    def add_cover_page(self):
        """Aggiungi la copertina di progetto."""

        cover_page = self.xml_root.find("design-cover-page")

        if cover_page is not None:
            # Big Titolo
            title = cover_page.find("h0").text if cover_page.find("h0") is not None else ""
            self.add_styled_paragraph(title, 'CoverH0', False, 1)
            self.add_empty_paragraphs(2)

            # Titolo
            title = cover_page.find("title").text if cover_page.find("title") is not None else ""
            self.add_styled_paragraph(title, 'CoverH1', False, 1)
            self.add_empty_paragraphs(1)

            # Sottotitolo
            subtitle = cover_page.find("subtitle")
            if subtitle is not None:
                self.add_empty_paragraphs(1)
                for text_element in subtitle.findall("text"):
                    self.add_styled_paragraph(text_element.text, style='CoverH2', upper=False, alignment=1)
            self.add_empty_paragraphs(1)

            # Sottotitolo
            today = cover_page.find("today")
            if today is not None:
                self.add_empty_paragraphs(1)
                self.add_styled_paragraph(today.text, style='Normal', upper=False, alignment=1)
            self.add_empty_paragraphs(1)

    def add_styled_paragraph(self, text, style, upper=False, alignment=2):
        """Aggiungi un paragrafo secondo lo stile indicato."""

        if text:
            if upper:
                text = text.upper()
            para = self.doc.add_paragraph(text, style=style)
            para.alignment = alignment

    def add_styled_table(self, table_node):
        if table_node is not None:
            word_table = None
            rows = table_node.findall("row")
            if rows:
                num_cols = len(rows[0])
                word_table = self.doc.add_table(rows=len(rows), cols=num_cols)
                for i, row in enumerate(rows):
                    cells = row.findall("cell")
                    for j, cell in enumerate(cells):
                        # Testo
                        # Usa uno spazio vuoto come valore predefinito
                        cell_text = cell.text if cell.text is not None else " "
                        word_table.cell(i, j).text = cell_text
                        # Immagine
                        image_node = cell.find("image")
                        if image_node is not None:
                            image_path = image_node.get("src")
                            if image_path:
                                image_data = self.storage.download_image(image_path)
                                paragraph = word_table.cell(i, j).paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(image_data)
                                paragraph.alignment = 1

            self.styles.apply_table_style(word_table, table_node.get('style'))

    def create_document(self):
        """Crea il documento Word a partire dall'XML."""

        self.add_metadata()
        self.add_cover_page()

        # Aggiungi il footer specifico solo per la copertina (contiene due tabelle revisione)
        first_section = self.doc.sections[0]
        # Abilita un piè di pagina specifico per la prima pagina
        first_section.footer.is_linked_to_previous = False
        first_section.different_first_page_header_footer = True

        # Sezione principale
        main_section = self.doc.add_section(WD_SECTION.NEW_PAGE)

        # Disabilita il collegamento per le intestazioni e i piè di pagina
        main_section.header.is_linked_to_previous = False
        main_section.footer.is_linked_to_previous = False

        # Inizia la numerazione della pagina in questa sezione a partire da 1
        main_section.start_page_number = 1  # QUESTO NON FUNZIONA

        # Aggiungi header e footer alla quarta pagina
        # self.add_header()
        self.add_footer(main_section)

        # Aggiungi i capitoli
        self.add_chapters()

    def save_document(self):
        """Salva il documento Word in un file temporaneo e lo carica nello storage."""
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = temp_file.name
                self.doc.save(temp_path)

            destination_path = self.storage.upload_file(temp_path, self.word_dest_path)
            logger.info(f"Documento salvato con successo in {destination_path}")
            return destination_path
        except Exception as e:
            logger.error(f"Errore nel salvataggio del documento nello storage: {e}", exc_info=True)
            return False
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    def create_and_save_word(self):
        """Gestisce l'intero processo: parsing XML, generazione Word, salvataggio."""

        # Converti stringa xml in root xml
        self.xml_root = self.convert_xml_string_in_node()

        if self.xml_root is None:
            logger.warning("Interrotto: XML non valido, impossibile generare il file Word.")
            return False

        # Crea il word
        self.create_document()

        # Salva il documento word
        return self.save_document()
